"""Safe text extraction from untrusted uploads.

Three separate ceilings apply, because one is not enough:

* the request byte cap, applied before this module is reached, bounds the
  upload;
* the archive guard bounds what a *small* upload is allowed to become once
  decompressed, since DOCX and XLSX are ZIP containers and a few kilobytes of
  repeated bytes can expand to gigabytes;
* the character accumulator bounds the extracted text as it is built, so a
  document that is legitimately huge is refused before it is held in memory
  rather than after.

Extraction runs in the background worker, never on the request path.
"""

import csv
import io
import json
import zipfile
from html.parser import HTMLParser
from pathlib import PurePath

SUPPORTED_DOCUMENT_EXTENSIONS = frozenset(
    {".txt", ".md", ".csv", ".tsv", ".json", ".html", ".htm", ".pdf", ".docx", ".xlsx"}
)
MAX_EXTRACTED_CHARACTERS = 2_000_000
DEFAULT_MAX_ARCHIVE_MEMBERS = 2_000
DEFAULT_MAX_EXPANSION_RATIO = 120

# A ceiling the expansion ratio cannot argue its way past: a large upload with a
# modest ratio is still capable of exhausting the worker.
_ABSOLUTE_ARCHIVE_BYTES = 300_000_000
_MAX_PDF_PAGES = 5_000
_MAX_WORKSHEETS = 200
_ZIP_ENCRYPTED_FLAG = 0x1


class DocumentExtractionError(ValueError):
    """A user-visible document cannot safely be converted to text."""


def is_supported_document(filename: str) -> bool:
    return PurePath(filename).suffix.lower() in SUPPORTED_DOCUMENT_EXTENSIONS


class _Accumulator:
    """Collects text and refuses to exceed the character budget.

    Raising mid-parse is the point: it stops a runaway spreadsheet at the cap
    instead of building the whole string and measuring it afterwards.
    """

    def __init__(self, limit: int) -> None:
        self._limit = limit
        self._length = 0
        self._parts: list[str] = []

    def add(self, text: str) -> None:
        if not text:
            return
        self._length += len(text) + 1
        if self._length > self._limit:
            raise DocumentExtractionError("The extracted document text is too large to index")
        self._parts.append(text)

    def text(self) -> str:
        return "\n".join(self._parts)


def extract_document_text(
    *,
    filename: str,
    content: bytes,
    max_characters: int = MAX_EXTRACTED_CHARACTERS,
    max_archive_members: int = DEFAULT_MAX_ARCHIVE_MEMBERS,
    max_expansion_ratio: int = DEFAULT_MAX_EXPANSION_RATIO,
) -> str:
    """Extract bounded text from one supported upload without trusting its MIME type."""
    extension = PurePath(filename).suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise DocumentExtractionError("This document format is not supported")

    if extension in {".docx", ".xlsx"}:
        _guard_archive(
            content,
            max_members=max_archive_members,
            max_expansion_ratio=max_expansion_ratio,
        )

    accumulator = _Accumulator(max_characters)
    try:
        if extension in {".txt", ".md"}:
            accumulator.add(content.decode("utf-8-sig"))
        elif extension in {".csv", ".tsv"}:
            _extract_delimited(
                content, delimiter="\t" if extension == ".tsv" else ",", into=accumulator
            )
        elif extension == ".json":
            accumulator.add(
                json.dumps(json.loads(content.decode("utf-8-sig")), ensure_ascii=False, indent=2)
            )
        elif extension in {".html", ".htm"}:
            _extract_html(content, into=accumulator)
        elif extension == ".pdf":
            _extract_pdf(content, into=accumulator)
        elif extension == ".docx":
            _extract_docx(content, into=accumulator)
        else:
            _extract_xlsx(content, into=accumulator)
    except DocumentExtractionError:
        raise
    except Exception as exc:
        # Parser-specific failures (including malformed PDFs and malformed
        # Office archives) are never exposed directly to API clients.
        raise DocumentExtractionError("The document could not be read as valid content") from exc

    normalized = "\n".join(
        line.strip() for line in accumulator.text().splitlines() if line.strip()
    )
    if not normalized:
        raise DocumentExtractionError("The document contains no readable text")
    return normalized


def _guard_archive(content: bytes, *, max_members: int, max_expansion_ratio: int) -> None:
    """Reject ZIP containers that are malformed, encrypted, or expansion bombs.

    This reads the central directory only. It never decompresses, so the check
    itself cannot be turned into the attack it is defending against.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            members = archive.infolist()
            if len(members) > max_members:
                raise DocumentExtractionError("The document contains too many internal parts")
            if any(member.flag_bits & _ZIP_ENCRYPTED_FLAG for member in members):
                raise DocumentExtractionError("Password-protected documents are not supported")
            uncompressed = sum(member.file_size for member in members)
    except zipfile.BadZipFile as exc:
        raise DocumentExtractionError("The document could not be read as valid content") from exc

    allowed = min(max(len(content), 1) * max_expansion_ratio, _ABSOLUTE_ARCHIVE_BYTES)
    if uncompressed > allowed:
        raise DocumentExtractionError("The document expands to an unsafe size")


def _extract_delimited(content: bytes, *, delimiter: str, into: _Accumulator) -> None:
    rows = csv.reader(io.StringIO(content.decode("utf-8-sig")), delimiter=delimiter)
    for row in rows:
        into.add(" | ".join(cell.strip() for cell in row))


class _TextCollector(HTMLParser):
    def __init__(self, into: _Accumulator) -> None:
        super().__init__()
        self._into = into

    def handle_data(self, data: str) -> None:
        self._into.add(data)


def _extract_html(content: bytes, *, into: _Accumulator) -> None:
    parser = _TextCollector(into)
    parser.feed(content.decode("utf-8-sig"))
    parser.close()


def _extract_pdf(content: bytes, *, into: _Accumulator) -> None:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    if reader.is_encrypted:
        raise DocumentExtractionError("Password-protected PDFs are not supported")
    if len(reader.pages) > _MAX_PDF_PAGES:
        raise DocumentExtractionError("The document has too many pages to index")
    for page in reader.pages:
        into.add(page.extract_text() or "")


def _extract_docx(content: bytes, *, into: _Accumulator) -> None:
    from docx import Document

    document = Document(io.BytesIO(content))
    for paragraph in document.paragraphs:
        into.add(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            into.add(" | ".join(cell.text for cell in row.cells))


def _extract_xlsx(content: bytes, *, into: _Accumulator) -> None:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    try:
        worksheets = workbook.worksheets
        if len(worksheets) > _MAX_WORKSHEETS:
            raise DocumentExtractionError("The workbook has too many sheets to index")
        for worksheet in worksheets:
            into.add(f"Sheet: {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    into.add(" | ".join(values))
    finally:
        workbook.close()
